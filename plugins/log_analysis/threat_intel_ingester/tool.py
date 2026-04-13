"""
Threat Intel Ingester — Event Mill Plugin

Ingests threat intelligence reports (PDF, HTML, STIX, CSV/JSON IOC lists)
and extracts structured IOC data with MITRE ATT&CK mapping.

This plugin depends on LLM capabilities for contextual IOC extraction
and MITRE technique inference. Regex-based extraction provides a baseline;
LLM analysis provides confidence scoring and priority assessment.
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger("eventmill.plugin.threat_intel_ingester")

# ---------------------------------------------------------------------------
# IOC Regex Patterns
# ---------------------------------------------------------------------------

IOC_PATTERNS = {
    "ip": re.compile(
        r"\b(?:(?:25[0-5]|2[0-4]\d|1?\d\d?)(?:\.|\[\.\]))"
        r"{3}(?:25[0-5]|2[0-4]\d|1?\d\d?)\b"
    ),
    "domain": re.compile(
        r"\b(?:[a-zA-Z0-9](?:[a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?"
        r"(?:\.|\[\.\]))+(?:com|net|org|io|info|biz|xyz|top|"
        r"ru|cn|de|uk|fr|jp|br|au|ca|nl|it|es|ch|se|no|fi|"
        r"dk|be|at|pl|cz|sk|hu|ro|bg|hr|si|lt|lv|ee|ie|pt|"
        r"gr|cy|lu|mt|li|is)\b",
        re.IGNORECASE,
    ),
    "hash_md5": re.compile(r"\b[a-fA-F0-9]{32}\b"),
    "hash_sha1": re.compile(r"\b[a-fA-F0-9]{40}\b"),
    "hash_sha256": re.compile(r"\b[a-fA-F0-9]{64}\b"),
    "url": re.compile(
        r"(?:https?|hxxps?|ftp)(?:://|(?:\[:\]//))[\w\-._~:/?#\[\]@!$&'()*+,;=%]+",
        re.IGNORECASE,
    ),
    "email": re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Z|a-z]{2,}\b"),
    "cve": re.compile(r"\bCVE-\d{4}-\d{4,7}\b", re.IGNORECASE),
    "mitre_technique": re.compile(r"\bT\d{4}(?:\.\d{3})?\b"),
}

# Defanging reversal patterns
DEFANG_REPLACEMENTS = [
    (re.compile(r"\[\.\]"), "."),
    (re.compile(r"hxxp", re.IGNORECASE), "http"),
    (re.compile(r"\[:\]"), ":"),
    (re.compile(r"\[at\]", re.IGNORECASE), "@"),
]


def refang(value: str) -> str:
    """Reverse common defanging patterns."""
    result = value
    for pattern, replacement in DEFANG_REPLACEMENTS:
        result = pattern.sub(replacement, result)
    return result


def was_defanged(original: str, refanged: str) -> bool:
    """Check if the value was defanged in the original text."""
    return original != refanged


# ---------------------------------------------------------------------------
# Text Extraction Helpers
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_path: str, max_pages: int = 50) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber is required for PDF processing")

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            if i >= max_pages:
                break
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_html(file_path: str) -> str:
    """Extract text from an HTML file using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError("beautifulsoup4 is required for HTML processing")

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    # Remove script and style elements
    for element in soup(["script", "style", "nav", "footer", "header"]):
        element.decompose()

    return soup.get_text(separator="\n", strip=True)


def extract_text_from_text(file_path: str) -> str:
    """Read a plain text file (including Markdown)."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for Word document processing. "
            "Install with: pip install python-docx"
        )
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


TEXT_EXTRACTORS = {
    "pdf_report": extract_text_from_pdf,
    "html_report": extract_text_from_html,
    "text": extract_text_from_text,
    "docx_report": extract_text_from_docx,
}


# ---------------------------------------------------------------------------
# Regex-Based IOC Extraction
# ---------------------------------------------------------------------------


@dataclass
class RawIOC:
    """An IOC extracted by regex before LLM refinement."""
    ioc_type: str
    value: str
    raw_value: str
    defanged: bool
    context: str = ""
    confidence: str = "low"
    priority: str = "medium"
    related_mitre: list[str] = field(default_factory=list)


def extract_iocs_regex(
    text: str,
    ioc_types: list[str],
) -> list[RawIOC]:
    """Extract IOCs from text using regex patterns.

    Returns deduplicated IOCs with surrounding context.
    """
    seen: set[tuple[str, str]] = set()
    results: list[RawIOC] = []

    for ioc_type in ioc_types:
        pattern = IOC_PATTERNS.get(ioc_type)
        if not pattern:
            continue

        for match in pattern.finditer(text):
            raw_value = match.group(0)
            value = refang(raw_value)
            defanged = was_defanged(raw_value, value)

            key = (ioc_type, value.lower())
            if key in seen:
                continue
            seen.add(key)

            # Extract surrounding context (up to 150 chars each side)
            start = max(0, match.start() - 150)
            end = min(len(text), match.end() + 150)
            context = text[start:end].replace("\n", " ").strip()

            results.append(
                RawIOC(
                    ioc_type=ioc_type,
                    value=value,
                    raw_value=raw_value,
                    defanged=defanged,
                    context=context[:300],
                    confidence="low",
                )
            )

    return results


# ---------------------------------------------------------------------------
# LLM Refinement Prompt
# ---------------------------------------------------------------------------

LLM_REFINEMENT_PROMPT = """You are an experienced threat intelligence analyst reviewing extracted IOCs from a security report.

TASK: Review the following IOC candidates and their surrounding context. For each IOC:
1. Assess whether it is a true indicator of compromise (not a benign version number, documentation example, or false positive)
2. Assign a confidence level: "low" (uncertain), "medium" (likely real IOC), "high" (confirmed IOC based on context)
3. Assign an operational priority: "low", "medium", "high" based on the IOC's role in the described attack
4. Identify related MITRE ATT&CK technique IDs if the context suggests a specific technique

Also identify any MITRE ATT&CK techniques described in the report text that are not already captured as IOC-type extractions. For each technique, note whether it was explicitly mentioned (technique ID appears in text) or inferred from described behavior.

SOURCE CONTEXT: {source_context}

IOC CANDIDATES:
{ioc_candidates}

FULL REPORT TEXT (truncated):
{report_text}

Respond ONLY with a JSON object in this exact format:
{{
  "refined_iocs": [
    {{
      "value": "the IOC value",
      "ioc_type": "ip|domain|hash_sha256|etc",
      "confidence": "low|medium|high",
      "priority": "low|medium|high",
      "context": "brief description of the IOC's role",
      "related_mitre": ["T1234", "T1234.001"],
      "is_false_positive": false
    }}
  ],
  "additional_mitre_techniques": [
    {{
      "technique_id": "T1234.001",
      "technique_name": "Technique Name",
      "tactic": "Tactic Name",
      "confidence": "explicit|inferred",
      "report_context": "brief description of the behavior"
    }}
  ],
  "report_metadata": {{
    "title": "report title if identifiable",
    "campaign_name": "named campaign if mentioned",
    "attributed_actor": "threat actor if attributed",
    "attribution_confidence": "low|medium|high"
  }}
}}
"""


# ---------------------------------------------------------------------------
# Protocol Types
# ---------------------------------------------------------------------------


@dataclass
class ToolResult:
    ok: bool
    result: dict[str, Any] | None = None
    error_code: str | None = None
    message: str | None = None
    details: dict[str, Any] | None = None
    output_artifacts: list[dict[str, Any]] | None = None


@dataclass
class ValidationResult:
    ok: bool
    errors: list[str] | None = None


# ---------------------------------------------------------------------------
# Tool Implementation
# ---------------------------------------------------------------------------


class ThreatIntelIngester:
    """Event Mill plugin: Threat Intelligence Report Ingester.

    Extracts IOCs from threat intelligence reports (PDF, HTML, text)
    using a two-pass approach:
    1. Regex-based extraction for baseline IOC identification
    2. LLM-based refinement for confidence scoring, false positive
       filtering, MITRE ATT&CK mapping, and priority assessment
    """

    def __init__(self) -> None:
        self._manifest: dict | None = None

    def _load_manifest(self) -> dict:
        if self._manifest is None:
            manifest_path = Path(__file__).parent / "manifest.json"
            with open(manifest_path) as f:
                self._manifest = json.load(f)
        return self._manifest

    def metadata(self) -> dict:
        """Return runtime metadata reflecting the manifest."""
        manifest = self._load_manifest()
        return {
            "tool_name": manifest["tool_name"],
            "version": manifest["version"],
            "pillar": manifest["pillar"],
            "display_name": manifest["display_name"],
            "description_short": manifest["description_short"],
            "stability": manifest["stability"],
            "requires_llm": manifest["requires_llm"],
            "artifacts_consumed": manifest["artifacts_consumed"],
            "artifacts_produced": manifest["artifacts_produced"],
        }

    def validate_inputs(self, payload: dict) -> ValidationResult:
        """Validate the input payload against the input schema."""
        errors = []

        if "artifact_id" not in payload:
            errors.append("artifact_id is required")

        if "ioc_types" in payload:
            valid_types = {
                "ip", "domain", "hash_md5", "hash_sha1", "hash_sha256",
                "url", "email", "cve", "mitre_technique",
            }
            for t in payload["ioc_types"]:
                if t not in valid_types:
                    errors.append(f"Unknown ioc_type: {t}")

        if "confidence_threshold" in payload:
            if payload["confidence_threshold"] not in ("low", "medium", "high"):
                errors.append("confidence_threshold must be low, medium, or high")

        if "max_pages" in payload:
            mp = payload["max_pages"]
            if not isinstance(mp, int) or mp < 1 or mp > 200:
                errors.append("max_pages must be an integer between 1 and 200")

        return ValidationResult(ok=len(errors) == 0, errors=errors if errors else None)

    def execute(self, payload: dict, context: Any) -> ToolResult:
        """Ingest a threat intelligence report and extract structured IOC data.

        Two-pass extraction:
        1. Regex pass: identify IOC candidates from raw text
        2. LLM pass: refine confidence, filter false positives, map MITRE techniques
        """
        artifact_id = payload.get("artifact_id")
        if not artifact_id:
            return ToolResult(
                ok=False,
                error_code="INPUT_VALIDATION_FAILED",
                message=(
                    "artifact_id is required. "
                    "Usage: run threat_intel_ingester {\"artifact_id\": \"<id>\"}"
                ),
            )
        source_context = payload.get("source_context", "")
        ioc_types = payload.get(
            "ioc_types",
            ["ip", "domain", "hash_sha256", "url", "cve", "mitre_technique"],
        )
        confidence_threshold = payload.get("confidence_threshold", "low")
        max_pages = payload.get("max_pages", 50)

        # --- Resolve artifact ---
        artifact = None
        for art in context.artifacts:
            if art.artifact_id == artifact_id:
                artifact = art
                break

        if artifact is None:
            return ToolResult(
                ok=False,
                error_code="ARTIFACT_NOT_FOUND",
                message=(
                    f"Artifact {artifact_id!r} not found in session. "
                    f"Use 'artifacts' to list loaded artifacts."
                ),
            )

        if artifact.artifact_type not in ("pdf_report", "html_report", "text", "docx_report"):
            return ToolResult(
                ok=False,
                error_code="INPUT_VALIDATION_FAILED",
                message=(
                    f"Artifact type '{artifact.artifact_type}' is not supported. "
                    f"Expected pdf_report, html_report, text (including .md), or docx_report."
                ),
            )

        # --- Extract text ---
        logger.info(
            "Extracting text from %s artifact %s",
            artifact.artifact_type,
            artifact_id,
        )
        extractor = TEXT_EXTRACTORS[artifact.artifact_type]
        try:
            if artifact.artifact_type == "pdf_report":
                raw_text = extractor(artifact.file_path, max_pages)
            else:
                raw_text = extractor(str(artifact.file_path))
        except Exception as e:
            logger.error("Text extraction failed: %s", e)
            return ToolResult(
                ok=False,
                error_code="ARTIFACT_UNREADABLE",
                message=f"Failed to extract text: {e}",
            )

        if artifact.artifact_type == "pdf_report":
            page_count = raw_text.count("\n\n") + 1
        else:
            page_count = len(raw_text.splitlines())

        # --- Regex extraction pass ---
        logger.info("Running regex IOC extraction for types: %s", ioc_types)
        raw_iocs = extract_iocs_regex(raw_text, ioc_types)
        logger.info("Regex pass found %d IOC candidates", len(raw_iocs))

        # --- LLM refinement pass ---
        refined_iocs = []
        mitre_mappings = []
        report_meta = {}

        if context.llm_enabled and context.llm_query is not None:
            logger.info("Running LLM refinement on %d IOC candidates", len(raw_iocs))

            ioc_candidates_text = "\n".join(
                f"- [{ioc.ioc_type}] {ioc.value} | Context: {ioc.context[:200]}"
                for ioc in raw_iocs[:100]  # Limit to avoid context overflow
            )

            prompt = LLM_REFINEMENT_PROMPT.format(
                source_context=source_context or "Not provided",
                ioc_candidates=ioc_candidates_text,
                report_text=raw_text[:8000],  # Truncate for context budget
            )

            # Use framework reference data for MITRE grounding
            grounding = []
            if hasattr(context, "reference_data"):
                mitre_data = context.reference_data.get("mitre_attack_enterprise")
                if mitre_data:
                    grounding.append(
                        "MITRE ATT&CK Enterprise techniques are available "
                        "for validation. Use official technique IDs."
                    )

            try:
                llm_response = context.llm_query.query_text(
                    prompt=prompt,
                    system_context=(
                        "You are a threat intelligence analyst. "
                        "Respond only with valid JSON."
                    ),
                    max_tokens=4096,
                    grounding_data=grounding,
                )

                if llm_response.ok and llm_response.text:
                    # Strip markdown code fences (LLMs often wrap JSON in ```json...```)
                    raw = llm_response.text.strip()
                    if raw.startswith("```"):
                        raw = re.sub(r"^```(?:json)?\s*\n?", "", raw)
                        raw = re.sub(r"\n?```\s*$", "", raw)
                    # Parse LLM JSON response
                    llm_result = json.loads(raw.strip())

                    # Build refined IOC list
                    for refined in llm_result.get("refined_iocs", []):
                        if refined.get("is_false_positive", False):
                            continue
                        refined_iocs.append(refined)

                    mitre_mappings = llm_result.get(
                        "additional_mitre_techniques", []
                    )
                    report_meta = llm_result.get("report_metadata", {})

                else:
                    logger.warning(
                        "LLM refinement failed, falling back to regex-only results. "
                        "Error: %s",
                        llm_response.error,
                    )
                    # Fall through to regex-only path

            except (json.JSONDecodeError, KeyError, TypeError) as e:
                logger.warning("LLM response parsing failed: %s", e)
                # Fall through to regex-only path

        # If LLM refinement didn't produce results, use regex baseline
        if not refined_iocs:
            logger.info("Using regex-only IOC results (no LLM refinement)")
            refined_iocs = [
                {
                    "ioc_type": ioc.ioc_type,
                    "value": ioc.value,
                    "confidence": "low",
                    "priority": "medium",
                    "context": ioc.context[:300],
                    "related_mitre": [],
                    "defanged": ioc.defanged,
                }
                for ioc in raw_iocs
            ]

        # --- Apply confidence threshold filter ---
        confidence_order = {"low": 0, "medium": 1, "high": 2}
        threshold_value = confidence_order.get(confidence_threshold, 0)
        filtered_iocs = [
            ioc
            for ioc in refined_iocs
            if confidence_order.get(ioc.get("confidence", "low"), 0) >= threshold_value
        ]

        # --- Build MITRE mappings from IOCs + additional techniques ---
        all_mitre = list(mitre_mappings)  # Start with additional techniques
        seen_techniques = {m["technique_id"] for m in all_mitre}

        for ioc in filtered_iocs:
            for tech_id in ioc.get("related_mitre", []):
                if tech_id not in seen_techniques:
                    seen_techniques.add(tech_id)
                    all_mitre.append(
                        {
                            "technique_id": tech_id,
                            "technique_name": "",  # Would be resolved from reference data
                            "tactic": "",
                            "confidence": "inferred",
                            "report_context": f"Associated with IOC {ioc['value']}",
                        }
                    )

        # --- Build summary ---
        ioc_breakdown: dict[str, int] = {}
        high_priority_count = 0
        confidence_dist = {"low": 0, "medium": 0, "high": 0}

        for ioc in filtered_iocs:
            ioc_type = ioc["ioc_type"]
            ioc_breakdown[ioc_type] = ioc_breakdown.get(ioc_type, 0) + 1
            if ioc.get("priority") == "high":
                high_priority_count += 1
            conf = ioc.get("confidence", "low")
            confidence_dist[conf] = confidence_dist.get(conf, 0) + 1

        # --- Register output artifact ---
        output_artifact_path = None
        output_artifact_id = None

        if hasattr(context, "register_artifact") and context.register_artifact:
            output_data = {
                "report_metadata": report_meta,
                "iocs": filtered_iocs,
                "mitre_mappings": all_mitre,
            }

            # Write artifact file
            import tempfile
            import os

            workspace = os.environ.get("EVENTMILL_WORKSPACE", "/tmp")
            artifact_dir = os.path.join(workspace, "artifacts")
            os.makedirs(artifact_dir, exist_ok=True)

            with tempfile.NamedTemporaryFile(
                mode="w",
                suffix="_ti_iocs.json",
                dir=artifact_dir,
                delete=False,
                prefix="art_",
            ) as f:
                json.dump(output_data, f, indent=2)
                output_artifact_path = f.name

            art_ref = context.register_artifact(
                artifact_type="json_events",
                file_path=output_artifact_path,
                source_tool="threat_intel_ingester",
                metadata={
                    "ioc_count": len(filtered_iocs),
                    "source_artifact": artifact_id,
                },
            )
            output_artifact_id = art_ref.artifact_id

        # --- Build result ---
        output_artifacts_list = None
        if output_artifact_id:
            output_artifacts_list = [
                {
                    "artifact_id": output_artifact_id,
                    "artifact_type": "json_events",
                    "file_path": output_artifact_path,
                    "description": "Structured IOC records extracted from threat intel report.",
                }
            ]

        logger.info(
            "Ingestion complete: %d IOCs, %d MITRE techniques, %d high-priority",
            len(filtered_iocs),
            len(all_mitre),
            high_priority_count,
        )

        return ToolResult(
            ok=True,
            result={
                "report_metadata": {
                    "title": report_meta.get("title", ""),
                    "source_organization": report_meta.get("source_organization", ""),
                    "publication_date": report_meta.get("publication_date", ""),
                    "page_count": page_count,
                    "artifact_type": artifact.artifact_type,
                    "campaign_name": report_meta.get("campaign_name", ""),
                    "attributed_actor": report_meta.get("attributed_actor", ""),
                    "attribution_confidence": report_meta.get(
                        "attribution_confidence", ""
                    ),
                },
                "iocs": filtered_iocs,
                "mitre_mappings": all_mitre,
                "summary": {
                    "total_iocs": len(filtered_iocs),
                    "ioc_breakdown": ioc_breakdown,
                    "high_priority_count": high_priority_count,
                    "mitre_technique_count": len(all_mitre),
                    "confidence_distribution": confidence_dist,
                },
            },
            output_artifacts=output_artifacts_list,
        )

    def summarize_for_llm(self, result: Any) -> str:
        """Produce a compressed summary for LLM context window."""
        if not result.ok:
            error = result.message or "Unknown error"
            return f"Threat intel ingestion failed: {error}"

        r = result.result or {}
        meta = r.get("report_metadata", {})
        summary = r.get("summary", {})
        mitre = r.get("mitre_mappings", [])

        parts = []

        # Report identity
        title = meta.get("title", "Unknown report")
        artifact_type = meta.get("artifact_type", "unknown")
        pages = meta.get("page_count", "?")
        size_label = "pages" if artifact_type == "pdf_report" else "lines"
        parts.append(f"Ingested {artifact_type} ({pages} {size_label}): {title}.")

        # Attribution
        actor = meta.get("attributed_actor")
        campaign = meta.get("campaign_name")
        if actor:
            conf = meta.get("attribution_confidence", "")
            parts.append(
                f"Attributed to {actor}"
                + (f" ({conf} confidence)" if conf else "")
                + (f", campaign: {campaign}" if campaign else "")
                + "."
            )

        # IOC counts
        total = summary.get("total_iocs", 0)
        breakdown = summary.get("ioc_breakdown", {})
        breakdown_str = ", ".join(
            f"{count} {ioc_type}{'s' if count != 1 else ''}"
            for ioc_type, count in sorted(breakdown.items())
        )
        parts.append(f"Extracted {total} IOCs: {breakdown_str}.")

        # High priority
        hp = summary.get("high_priority_count", 0)
        if hp > 0:
            parts.append(f"{hp} IOCs flagged as high-priority.")

        # MITRE techniques
        tech_count = summary.get("mitre_technique_count", 0)
        if tech_count > 0 and mitre:
            tech_list = ", ".join(
                f"{m['technique_id']} ({m.get('technique_name', '')})"
                for m in mitre[:6]
            )
            parts.append(f"Mapped to {tech_count} MITRE techniques: {tech_list}.")
            if tech_count > 6:
                parts.append(f"(and {tech_count - 6} more)")

        # Output artifact
        artifacts = result.output_artifacts or []
        if artifacts:
            art = artifacts[0]
            parts.append(
                f"Output artifact: {art['artifact_id']} "
                f"({art['artifact_type']})."
            )

        return " ".join(parts)
